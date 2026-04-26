"""
Generate the EMBRACE AI — 12-Month Initiative Report (Word Document)
Run this script once to produce the .docx file.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "EMBRACE_AI_12_Month_Initiative_Report.docx")

doc = Document()

# ─── Styles & Colours ───
PETROL = RGBColor(0x00, 0x99, 0x99)
DARK = RGBColor(0x00, 0x00, 0x28)
GREY = RGBColor(0x6E, 0x6E, 0x73)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF7, 0xF7, 0xF8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = PETROL if level <= 2 else DARK
    hs.font.name = 'Calibri'

# ─── Helpers ───
def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_para("SIEMENS", bold=True, size=28, color=PETROL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Engineering Systems", bold=False, size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()
add_para("EMBRACE AI", bold=True, size=36, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para("12-Month AI Initiative Plan", bold=False, size=18, color=PETROL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Engineering Systems · FT D AA IN SGI DET ENGSYS", size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Siemens, Chennai, India", size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

doc.add_paragraph()
doc.add_paragraph()
add_para("Prepared by: Vigneshvar SA", size=12, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para("vigneshvar.sa@siemens.com", size=11, color=PETROL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("April 2025", size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. About Engineering Systems",
    "3. Initiative Overview",
    "4. Team Members",
    "5. Syllabus Design Philosophy",
    "6. Points & Rewards System",
    "7. Recurring Monthly Structure",
    "8. Seminar Table Topics",
    "9. Month-by-Month Detailed Plan",
    "   9.1  Month 1 — Who's Who: AI Edition",
    "   9.2  Month 2 — AI Storyteller: The EngSys Chronicles",
    "   9.3  Month 3 — Prompt Battle Arena",
    "   9.4  Month 4 — My AI Assistant: Copilot Power User",
    "   9.5  Month 5 — Code with AI",
    "   9.6  Month 6 — AI Knowledge Base: Smart Documentation",
    "   9.7  Month 7 — Automate the Boring Stuff",
    "   9.8  Month 8 — AI Agents Hackathon",
    "   9.9  Month 9 — AI + Data: Data Detective",
    "   9.10 Month 10 — My AI Use Case: Problem Definition",
    "   9.11 Month 11 — My AI Use Case: Build & Iterate",
    "   9.12 Month 12 — Grand Finale Showcase",
    "10. EMBRACE AI Dashboard (Web Application)",
    "11. Expected Outcomes",
    "12. Cumulative Deliverables",
    "13. Appendix — Event Calendar",
]
for item in toc_items:
    add_para(item, size=11, space_after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "EMBRACE AI is a structured, gamified, 12-month initiative designed to progressively upskill "
    "the 16-member Engineering Systems offshore team at Siemens, Chennai on Generative AI. "
    "The initiative transitions from fun, curiosity-driven activities in the early months to "
    "hands-on automation, agentic AI applications, and ultimately the delivery of individual "
    "business use-cases demonstrating measurable value enabled by AI."
)
doc.add_paragraph(
    "The initiative is supported by a custom-built web dashboard that provides a live leaderboard, "
    "event calendar, real-time quiz platform, anonymous seminar survey system, and a comprehensive "
    "admin panel — all accessible to the entire team via a single URL."
)
doc.add_paragraph(
    "By the end of 12 months, the team will have: (a) developed fluency in GenAI tools used at "
    "Siemens, (b) built 16 documented AI use-case prototypes aligned to PLM, Manufacturing, "
    "and OT domains, and (c) created a replicable playbook that can be shared across other "
    "Siemens teams in India."
)

# ═══════════════════════════════════════════════════════════
# 2. ABOUT ENGINEERING SYSTEMS
# ═══════════════════════════════════════════════════════════
doc.add_heading("2. About Engineering Systems", level=1)
doc.add_paragraph(
    "The Engineering Systems (EngSys) team serves as the digital and operational backbone for "
    "partner organizations. The team fully manages, maintains, and customizes critical applications "
    "and infrastructure required for both Engineering and Manufacturing operations."
)

doc.add_heading("Three Pillars of Expertise", level=2)

doc.add_heading("Pillar 1: PLM (Product Lifecycle Management)", level=3)
doc.add_paragraph(
    "Champions in Teamcenter Administration & Customization, MCAD Support (Creo and NX Open), "
    "ECAD Support (Zuken e3 / Panel Builder), and SAP Integration (T4S). "
    "Actively engaged in Rulestream Customization and Support."
)

doc.add_heading("Pillar 2: Manufacturing", level=3)
doc.add_paragraph(
    "Champions in OpCenter ED Management (MES), OpCenter OEE Management, and Praxis Management "
    "(Nesting Programs). Actively engaged in AGW Administration and DMM Administration."
)

doc.add_heading("Pillar 3: OT & Automation (Operational Technology)", level=3)
doc.add_paragraph(
    "Champions in OT Network Management (SINEC NMS) and Data Analytics. Actively engaged in "
    "SCADA & PLC System Administration, Industrial Information Hub, Industrial Edge, and "
    "Artificial Intelligence (GenAI). Future engagements include AVR/AMR Solutions."
)

doc.add_heading("Cross-Functional Support", level=3)
doc.add_paragraph(
    "Infrastructure Management, Application Lifecycle (Greenfield & Brownfield), and Project Management."
)

# ═══════════════════════════════════════════════════════════
# 3. INITIATIVE OVERVIEW
# ═══════════════════════════════════════════════════════════
doc.add_heading("3. Initiative Overview", level=1)

add_table(
    ["Attribute", "Details"],
    [
        ["Initiative Name", "EMBRACE AI"],
        ["Team", "Engineering Systems (EngSys)"],
        ["Location", "Siemens, Chennai, India (Offshore)"],
        ["Reporting To", "North America Profit Centers"],
        ["Team Size", "16 members"],
        ["Duration", "12 months (April 2025 – March 2026)"],
        ["Initiative Lead", "Vigneshvar SA"],
        ["Cadence", "1 session per month (2nd Tuesday)"],
        ["Session Duration", "90–150 minutes"],
        ["Prizes", "Star Points → Gift Cards / Vouchers for Top 3"],
    ],
)

# ═══════════════════════════════════════════════════════════
# 4. TEAM MEMBERS
# ═══════════════════════════════════════════════════════════
doc.add_heading("4. Team Members", level=1)

members = [
    ["Aravinthan Dhanabal", "aravinthan.dhanabal@siemens.com", "Manufacturing"],
    ["Aswin Gangadharan", "aswin.gangadharan@siemens.com", "OT & Automation"],
    ["Balaji M", "balaji.m@siemens.com", "PLM"],
    ["Dharankumar N", "dharankumar.n@siemens.com", "PLM"],
    ["Gopalakrishnan Rajendran", "gopalakrishnan.rajendran@siemens.com", "PLM"],
    ["John Yabesh Johnson", "john-yabesh.johnson@siemens.com", "Manufacturing"],
    ["Minal Ramesh Nilange", "minal.nilange@siemens.com", "Manufacturing"],
    ["Pandiyarajan Nagarajan", "pandiyarajan.n@siemens.com", "Cross-Functional"],
    ["Priyadharshani Varsha S", "priyadharshani.varsha-s@siemens.com", "PLM"],
    ["Rizwan Ali M", "rizwan.m@siemens.com", "PLM"],
    ["Saravanan Jayabalan", "saravanan.jayabalan@siemens.com", "Cross-Functional"],
    ["Sivakumar D", "sivakumar.d.ext@siemens.com", "Cross-Functional"],
    ["Soniya Dhayalan", "soniya.dhayalan@siemens.com", "OT & Automation"],
    ["Tippu Sultan Shaik", "tippu.shaik@siemens.com", "PLM"],
    ["Vidya Srivatsan", "vidya.srivatsan@siemens.com", "Cross-Functional"],
    ["Vigneshvar SA (Lead)", "vigneshvar.sa@siemens.com", "PLM"],
]

add_table(
    ["#", "Name", "Email", "Domain"],
    [[i+1, *m] for i, m in enumerate(members)],
)

# ═══════════════════════════════════════════════════════════
# 5. SYLLABUS DESIGN PHILOSOPHY
# ═══════════════════════════════════════════════════════════
doc.add_heading("5. Syllabus Design Philosophy", level=1)

doc.add_paragraph(
    "The syllabus follows a four-phase maturity model, transitioning from fun engagement "
    "to measurable business impact:"
)

add_table(
    ["Phase", "Months", "Theme", "Orientation"],
    [
        ["SPARK", "1–3", "Curiosity & Fun", "80% Fun / 20% Learning"],
        ["BUILD", "4–6", "Skill Building & Tooling", "50% Fun / 50% Learning"],
        ["APPLY", "7–9", "Domain-Specific Automation", "20% Fun / 80% Business"],
        ["DELIVER", "10–12", "Business Value Demonstration", "100% Business Impact"],
    ],
)

doc.add_paragraph(
    "Each phase builds on the previous one. By the time members reach Phase 4, they will "
    "have accumulated enough skills and confidence from Phases 1–3 to independently design, "
    "build, and present an AI-enabled business use-case."
)

# ═══════════════════════════════════════════════════════════
# 6. POINTS & REWARDS SYSTEM
# ═══════════════════════════════════════════════════════════
doc.add_heading("6. Points & Rewards System", level=1)

doc.add_paragraph(
    "A gamified points system runs throughout the 12 months. Points are accumulated across "
    "all activities and tracked on the EMBRACE AI Dashboard leaderboard."
)

add_table(
    ["Category", "Points"],
    [
        ["Seminar Presenter (peer-rated 1–5)", "Up to 5 pts"],
        ["Quiz Winner — 1st Place", "3 pts"],
        ["Quiz Winner — 2nd Place", "2 pts"],
        ["Quiz Winner — 3rd Place", "1 pt"],
        ["Activity Completion (on time)", "3 pts"],
        ["Activity Category Awards", "5 pts each"],
        ["First to Submit (Fast + Quality)", "2 pts bonus"],
        ["People's Choice (peer vote)", "3 pts"],
        ["Grand Finale — Best Overall Use Case", "15 pts"],
        ["Grand Finale — Runner-Up", "10 pts"],
        ["Grand Finale — Second Runner-Up", "7 pts"],
    ],
)

doc.add_paragraph(
    "Year-End Prizes: The top 3 cumulative point-earners receive Star Points, which can "
    "be redeemed for gift cards, coupons, or vouchers. All 16 members receive certificates "
    "of completion."
)

# ═══════════════════════════════════════════════════════════
# 7. RECURRING MONTHLY STRUCTURE
# ═══════════════════════════════════════════════════════════
doc.add_heading("7. Recurring Monthly Structure", level=1)

doc.add_paragraph(
    "Every monthly session follows a consistent structure to build rhythm and expectation:"
)

add_table(
    ["Segment", "Duration", "Description"],
    [
        ["Seminar (Table Topic)", "15–20 min", "1 randomly assigned person presents a topic drawn from the bowl the previous month. Followed by a 5-question quiz for the audience."],
        ["Main Activity", "45–90 min", "Either spontaneous (in-session) or pre-worked (submitted before session)."],
        ["Leaderboard Update & Awards", "10 min", "Points tallied, leaderboard updated on the dashboard, category winners announced."],
    ],
)

# ═══════════════════════════════════════════════════════════
# 8. SEMINAR TABLE TOPICS
# ═══════════════════════════════════════════════════════════
doc.add_heading("8. Seminar Table Topics (Draw from Bowl)", level=1)

doc.add_paragraph(
    "Twelve topics are prepared on paper slips and placed in a physical bowl. At the end "
    "of each session, one person draws a slip for the following month. Month 1 and Month 12 "
    "seminars are presented by Vigneshvar SA."
)

topics = [
    ["1", "Vigneshvar SA (fixed)", "What is Generative AI and Why Should Engineers Care?"],
    ["2", "Random draw", "History & Evolution of AI — From Turing to Transformers"],
    ["3", "Random draw", "Prompt Engineering Deep Dive — Techniques, Patterns & Anti-Patterns"],
    ["4", "Random draw", "M365 Copilot — What It Can Actually Do for Engineers"],
    ["5", "Random draw", "AI-Assisted Coding — GitHub Copilot, SiemensGPT, Best Practices"],
    ["6", "Random draw", "RAG — How AI Can Search and Summarize Your Own Documents"],
    ["7", "Random draw", "Agentic AI — What Are AI Agents vs. Simple Chatbots"],
    ["8", "Random draw", "Building AI Agents — Tools, Frameworks, and Patterns"],
    ["9", "Random draw", "AI for Data Analysis — LLMs for Querying, Summarizing & Visualizing Data"],
    ["10", "Random draw", "AI in Industrial Engineering — Real-World Case Studies"],
    ["11", "Random draw", "Responsible AI — Ethics, Bias, Compliance & Siemens Principles"],
    ["12", "Vigneshvar SA (fixed)", "Our AI Journey: 12 Months of Growth (Retrospective)"],
]
add_table(["Month", "Presenter", "Topic"], topics)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 9. MONTH-BY-MONTH DETAILED PLAN
# ═══════════════════════════════════════════════════════════
doc.add_heading("9. Month-by-Month Detailed Plan", level=1)

# ─── Data for all 12 months ───
months_data = [
    {
        "num": 1, "title": "Who's Who: AI Edition", "phase": "SPARK", "date": "April 8, 2025",
        "duration": "90 min", "type": "Pre-worked (submitted 2 weeks before session)",
        "seminar": "What is Generative AI and Why Should Engineers Care? (Vigneshvar SA)",
        "description": (
            "The inaugural activity that officially kicks off EMBRACE AI. Each team member generates "
            "AI portraits and one-liners for every other colleague using GenAI tools."
        ),
        "prep": [
            "Send kickoff email with detailed instructions and reference links.",
            "Create Microsoft Form for submissions (Colleague Name, One-liner, Prompt, Image).",
            "Share prompt engineering resources: learnprompting.org, promptingguide.ai.",
            "Share image generation guides: OpenAI images guide, Leonardo.ai blog.",
            "Set submission deadline: 2 weeks from email date.",
            "Prepare 5-question quiz on GenAI basics.",
            "Prepare award certificates/badges for 4 categories.",
        ],
        "execution": [
            "0:00–0:20 — Seminar: Vigneshvar presents 'What is GenAI and Why Should Engineers Care?'",
            "0:20–0:30 — Quiz: 5 questions on GenAI basics. Top 3 scorers get 3/2/1 points.",
            "0:30–0:35 — Official kickoff: Unveil EMBRACE AI Charter, explain 12-month journey.",
            "0:35–1:00 — Gallery Walk: Display all images and one-liners on shared screen.",
            "1:00–1:15 — Peer Voting for each award category.",
            "1:15–1:25 — Award Ceremony. Update leaderboard live on the dashboard.",
            "1:25–1:30 — Draw from the bowl for next month's seminar presenter.",
        ],
        "awards": ["Most Creative Prompt", "Best AI-Generated Image", "Funniest One-Liner", "First to Complete"],
        "skills": "Prompt engineering basics, image generation, iterative prompting",
        "tools": "SiemensGPT, M365 Copilot, Gemini (with caution), DALL-E, Leonardo.ai",
        "followup": [
            "Update leaderboard with all Month 1 points.",
            "Post recap in Teams channel with session photos and winners.",
            "Remind next seminar presenter (drawn from bowl).",
            "Collect feedback: 3-question pulse check.",
        ],
    },
    {
        "num": 2, "title": "AI Storyteller: The EngSys Chronicles", "phase": "SPARK", "date": "May 13, 2025",
        "duration": "90 min", "type": "In-session (Groups of 4)",
        "seminar": "History & Evolution of AI — From Turing to Transformers",
        "description": (
            "Teams of 4 create AI-generated illustrated stories from fictional EngSys scenarios. "
            "Comic-strip style using AI-generated narration and images."
        ),
        "prep": [
            "Create 4 fictional EngSys scenario cards (Teamcenter crash, OEE anomaly, OT network mystery, Rulestream bug).",
            "Pre-assign 4 groups of 4 — mix across PLM, Manufacturing, OT, Cross-functional.",
            "Prepare 4–6 panel comic strip templates (PowerPoint or Canva).",
            "Prepare 5-question quiz on AI history.",
        ],
        "execution": [
            "0:00–0:20 — Seminar by drawn presenter.",
            "0:20–0:30 — Quiz on AI history.",
            "0:30–0:35 — Activity briefing. Distribute scenario cards.",
            "0:35–1:05 — Group work (30 min): Create comic-strip stories using AI.",
            "1:05–1:25 — Group presentations (5 min each).",
            "1:25–1:30 — Voting & awards. Draw next topic.",
        ],
        "awards": ["Best Storyline", "Most Visually Creative", "Best Use of AI Tools", "People's Choice"],
        "skills": "Multi-turn prompting, chaining prompts, collaborative AI use, visual storytelling",
        "tools": "SiemensGPT, M365 Copilot, image generators",
        "followup": [
            "Update leaderboard. Post best comic strips in Teams channel.",
            "Consider printing the best comic strip for the team area.",
        ],
    },
    {
        "num": 3, "title": "Prompt Battle Arena", "phase": "SPARK", "date": "June 10, 2025",
        "duration": "100 min", "type": "In-session (Individual tournament)",
        "seminar": "Prompt Engineering Deep Dive — Techniques, Patterns & Anti-Patterns",
        "description": (
            "A live prompt engineering tournament with Speed Round, Refinement Round, "
            "Domain Challenge, and Bonus Round."
        ),
        "prep": [
            "Design Round 1 (Speed): 1 vague task prompt for everyone.",
            "Design Round 2 (Refinement): 1 deliberately bad prompt to improve in 3 iterations.",
            "Design Round 3 (Domain): 16 domain-specific tasks matched to each member's pillar.",
            "Design Bonus Round: 1 mediocre prompt to improve with constraints only.",
            "Set up live display for real-time output comparison.",
            "Prepare scoring rubrics for each round.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on prompt engineering.",
            "0:20–0:30 — Quiz on prompting techniques.",
            "0:30–0:40 — Round 1 (Speed): 3 min per person, peer vote.",
            "0:40–0:55 — Round 2 (Refinement): 3 iterations, 5 min total.",
            "0:55–1:15 — Round 3 (Domain): 5 min per person, judged on relevance + quality.",
            "1:15–1:25 — Bonus Round: Constraints-only improvement.",
            "1:25–1:35 — Awards. Draw next topic.",
        ],
        "awards": ["Speed King/Queen", "Best Iterative Improver", "Domain Prompt Master", "Overall Prompt Champion"],
        "skills": "Advanced prompt engineering — personas, system instructions, few-shot, chain-of-thought",
        "tools": "SiemensGPT, M365 Copilot",
        "followup": [
            "Share winning prompts in Teams as a 'Prompt Playbook' reference document.",
            "Compile each person's best prompt into a shared 'EngSys Prompt Library'.",
        ],
    },
    {
        "num": 4, "title": "My AI Assistant: Copilot Power User", "phase": "BUILD", "date": "July 8, 2025",
        "duration": "100 min", "type": "Pre-worked (2 weeks) + Live demo",
        "seminar": "M365 Copilot — What It Can Actually Do for Engineers",
        "description": (
            "Each member picks 1 real recurring work task and uses M365 Copilot to automate or "
            "significantly accelerate it. Live 3-minute demos in session."
        ),
        "prep": [
            "Send assignment: Pick 1 recurring task, use M365 Copilot to automate it.",
            "Create submission form: Task Description, Prompt(s), Output, Time Saved, Screenshot.",
            "Share M365 Copilot tutorials and Copilot Lab prompts library.",
            "Set submission deadline 3 days before session.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on M365 Copilot ecosystem.",
            "0:20–0:30 — Quiz on Copilot capabilities.",
            "0:30–1:10 — Demo round: Each person gives 2–3 min live demo.",
            "1:10–1:20 — Group discussion: What worked, what didn't.",
            "1:20–1:30 — Awards. Draw next topic.",
        ],
        "awards": ["Most Time Saved", "Most Creative Use of Copilot", "Best Before/After Transformation", "People's Choice"],
        "skills": "Practical AI tool integration, measuring AI ROI (time saved), M365 Copilot deep dive",
        "tools": "M365 Copilot (Word, Excel, PowerPoint, Teams, Outlook)",
        "followup": [
            "Compile all 16 workflows into a 'Copilot Cookbook for EngSys' document.",
            "Calculate total aggregate time saved — share with manager as ROI metric.",
        ],
    },
    {
        "num": 5, "title": "Code with AI", "phase": "BUILD", "date": "August 12, 2025",
        "duration": "100 min", "type": "Pre-worked (2 weeks)",
        "seminar": "AI-Assisted Coding — GitHub Copilot, SiemensGPT for Code, Best Practices",
        "description": (
            "Domain-specific coding tasks using AI. Members annotate what was correct, incorrect, "
            "and hallucinated. Final working versions shared."
        ),
        "prep": [
            "Assign domain-specific coding tasks (Teamcenter ITK, NX Open, Creo Toolkit, PLC logic, SQL for OEE, Python for logs, PowerShell for health checks, etc.).",
            "Create submission form: Task, Prompt, AI Code, Annotations, Final Version.",
            "Share resources on AI code generation best practices and hallucination risks.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on AI-assisted coding.",
            "0:20–0:30 — Quiz on code generation safety.",
            "0:30–1:00 — Show & Tell: Each person walks through their code task (2 min each).",
            "1:00–1:15 — 'Hallucination Hall of Fame': Funniest/worst AI code hallucinations.",
            "1:15–1:25 — Key takeaways discussion.",
            "1:25–1:35 — Awards. Draw next topic.",
        ],
        "awards": ["Most Accurate AI-Generated Code", "Best Prompt for Code Generation", "Best Code Review / Bug Catch", "Most Relevant to Real Work"],
        "skills": "AI-assisted coding, code review discipline, understanding hallucinations",
        "tools": "SiemensGPT, M365 Copilot, GitHub Copilot (if available)",
        "followup": [
            "Add annotated code examples to 'Code with AI — Lessons Learned' document.",
            "Compile 'Hallucination Patterns' reference for each domain.",
        ],
    },
    {
        "num": 6, "title": "AI Knowledge Base: Smart Documentation", "phase": "BUILD", "date": "September 9, 2025",
        "duration": "100 min", "type": "Collaborative + In-session",
        "seminar": "RAG — How AI Can Search and Summarize Your Own Documents",
        "description": (
            "Team builds a prototype AI-powered Q&A system over team documentation. "
            "Each person tests with domain-specific questions and rates accuracy."
        ),
        "prep": [
            "Each person contributes 2–3 non-sensitive documents (SOPs, runbooks, config manuals).",
            "Collect documents into shared folder organized by pillar.",
            "Test document upload with SiemensGPT/Copilot file referencing.",
            "Prepare 3 sample questions per domain.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on RAG.",
            "0:20–0:30 — Quiz on RAG concepts.",
            "0:30–0:45 — Live Build: Upload documents, set up knowledge assistant.",
            "0:45–1:10 — Testing: Each person asks 3 questions, rates accuracy.",
            "1:10–1:20 — Discussion: What worked, what failed, what makes good source docs.",
            "1:20–1:30 — Awards. Draw next topic.",
        ],
        "awards": ["Best Source Documents Contributed", "Most Insightful Test Questions", "Best Improvement Suggestion", "RAG Champion"],
        "skills": "RAG fundamentals, document preparation for AI, enterprise AI assistants",
        "tools": "SiemensGPT (document upload), M365 Copilot",
        "followup": [
            "Publish 'Document Preparation Best Practices for AI' guide.",
            "Assign volunteers to maintain/improve the knowledge assistant.",
            "MID-YEAR CHECK-IN: Share leaderboard summary with manager, collect feedback, adjust plan.",
        ],
    },
    {
        "num": 7, "title": "Automate the Boring Stuff", "phase": "APPLY", "date": "October 14, 2025",
        "duration": "110 min", "type": "Pre-worked (3 weeks) + Live demo",
        "seminar": "Agentic AI — What Are AI Agents vs. Simple Chatbots",
        "description": (
            "Each person identifies and prototypes an AI automation for 1 repetitive daily task. "
            "5-minute presentations with live demos."
        ),
        "prep": [
            "Send assignment: Identify 1 repetitive task, document Problem/Solution/Prototype/Impact.",
            "Share examples by pillar (auto-generate change requests, summarize OEE reports, parse network logs, etc.).",
            "Offer 1:1 brainstorming slots (15-min optional calls).",
        ],
        "execution": [
            "0:00–0:20 — Seminar on Agentic AI.",
            "0:20–0:30 — Quiz.",
            "0:30–1:20 — Demo presentations: 5 min each (problem → solution → demo).",
            "1:20–1:35 — Peer voting and discussion.",
            "1:35–1:45 — Awards. Draw next topic.",
        ],
        "awards": ["Most Impactful Automation", "Most Creative Solution", "Best Prototype / Demo", "Highest Estimated Time Savings"],
        "skills": "Identifying automation opportunities, AI solution design, basic agentic thinking, ROI estimation",
        "tools": "SiemensGPT, M365 Copilot, Power Automate, Python",
        "followup": [
            "Compile 'Automation Catalog' — one-page summary of each automation.",
            "Calculate total aggregate time saved. Share with manager.",
        ],
    },
    {
        "num": 8, "title": "AI Agents Hackathon", "phase": "APPLY", "date": "November 11, 2025",
        "duration": "120 min", "type": "In-session (Groups — Extended session)",
        "seminar": "Building AI Agents — Tools, Frameworks, and Patterns",
        "description": (
            "Team hackathon with complex multi-step engineering scenarios. Groups design agent "
            "architectures on whiteboards and implement key steps using GenAI tools."
        ),
        "prep": [
            "Design 4 complex multi-step scenarios (PLM variant setup, OEE root cause analysis, SINEC NMS alarm flood, Teamcenter migration).",
            "Pre-assign 4 mixed-pillar groups of 4 (different from Month 2).",
            "Prepare whiteboard/flipchart supplies.",
            "Book extended 2-hour session.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on building AI agents.",
            "0:20–0:30 — Quiz.",
            "0:30–0:40 — Hackathon briefing. Distribute scenarios.",
            "0:40–1:20 — Hackathon (40 min): Design agent architecture + implement key steps.",
            "1:20–1:45 — Group presentations (6 min each).",
            "1:45–1:55 — Awards. Draw next topic.",
        ],
        "awards": ["Best Agent Architecture Design", "Most Realistic Implementation", "Best Team Collaboration", "Most Business-Relevant Solution"],
        "skills": "Multi-step AI reasoning, agent design patterns, tool orchestration, complex problem decomposition",
        "tools": "SiemensGPT, M365 Copilot, Whiteboard",
        "followup": [
            "Photograph all whiteboard architectures, share in Teams.",
            "Document best agent designs as 'Agent Design Patterns for EngSys' reference.",
        ],
    },
    {
        "num": 9, "title": "AI + Data: Data Detective", "phase": "APPLY", "date": "December 9, 2025",
        "duration": "100 min", "type": "Pre-worked (2–3 weeks) + Live presentation",
        "seminar": "AI for Data Analysis — LLMs for Querying, Summarizing & Visualizing Data",
        "description": (
            "Each person analyzes a domain-specific dataset using AI tools. Ask 5 natural-language "
            "questions, generate 1 visualization, identify 1 insight/anomaly."
        ),
        "prep": [
            "Prepare/source sample datasets per pillar (BOM data, OEE data, network alarm logs, ticket data).",
            "Distribute datasets. Send instructions.",
            "Create submission form.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on AI for data analysis.",
            "0:20–0:30 — Quiz.",
            "0:30–1:10 — Data Detective presentations (2.5 min each).",
            "1:10–1:20 — Discussion: Surprising insights, what worked, limitations.",
            "1:20–1:30 — Awards. Draw next topic.",
        ],
        "awards": ["Best Insight / Anomaly Discovery", "Best AI-Generated Visualization", "Most Business-Relevant Finding", "Data Detective Champion"],
        "skills": "AI-powered data analysis, natural language querying, data visualization with AI",
        "tools": "M365 Copilot (Excel), SiemensGPT, Python",
        "followup": [
            "Compile all insights into 'AI Data Analysis Findings' doc.",
            "Escalate any real business insights to the manager.",
            "END-OF-PHASE-3 CHECKPOINT: Share cumulative leaderboard. Announce capstone phase.",
        ],
    },
    {
        "num": 10, "title": "My AI Use Case: Problem Definition", "phase": "DELIVER", "date": "January 13, 2026",
        "duration": "110 min", "type": "Pre-worked (3 weeks) + In-session workshop",
        "seminar": "AI in Industrial Engineering — Real-World Case Studies",
        "description": (
            "Capstone kickoff. Each person identifies 1 real business use case from their own work area "
            "where AI can provide measurable value. Formal 1-page proposal."
        ),
        "prep": [
            "Send capstone kickoff email with use case proposal template (Problem, Current Process, Proposed AI Solution, Expected Impact, Feasibility).",
            "Offer 1:1 brainstorming sessions (15 min per person).",
            "Collect proposals 3 days before session.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on AI in industrial engineering.",
            "0:20–0:30 — Quiz.",
            "0:30–0:35 — Capstone briefing: 3-month journey (Define → Build → Showcase).",
            "0:35–1:20 — Use case pitches (3 min each). Team asks 1 question each.",
            "1:20–1:30 — Peer feedback. Assign buddy pairs for support.",
            "1:30–1:40 — Awards. Draw next topic.",
        ],
        "awards": ["Best Problem Definition", "Most Innovative Proposed Solution", "Most Impactful Use Case (peer-voted)", "Best Feasibility Analysis"],
        "skills": "Business problem framing, AI solution design, stakeholder communication, feasibility assessment",
        "tools": "All tools from previous months",
        "followup": [
            "Finalize all 16 use case proposals. Share with manager.",
            "Set up buddy pair check-in schedule (bi-weekly).",
            "Create shared progress tracker (Kanban/Planner).",
        ],
    },
    {
        "num": 11, "title": "My AI Use Case: Build & Iterate", "phase": "DELIVER", "date": "February 10, 2026",
        "duration": "120 min", "type": "Pre-worked (ongoing) + Peer review session",
        "seminar": "Responsible AI — Ethics, Bias, Compliance & Siemens Principles",
        "description": (
            "Progress review — each person demonstrates their working prototype. Round-robin "
            "peer reviews in groups of 4. Office hours for hands-on help."
        ),
        "prep": [
            "Weekly nudges: Vigneshvar sends Teams message every Monday.",
            "Buddy check-ins: Each pair meets twice during the month.",
            "Submission: Working prototype, demo video/screenshots, challenges documented.",
        ],
        "execution": [
            "0:00–0:20 — Seminar on Responsible AI.",
            "0:20–0:30 — Quiz.",
            "0:30–1:10 — Round-robin reviews (groups of 4): Demo, structured feedback.",
            "1:10–1:30 — Office hours: Hands-on help for anyone stuck.",
            "1:30–1:40 — Iteration planning: 3 improvement items each.",
            "1:40–1:50 — Awards.",
        ],
        "awards": ["Most Progress Made", "Best Working Prototype", "Best Peer Feedback Given", "Most Resilient"],
        "skills": "Iterative development, peer review, giving/receiving technical feedback, responsible AI",
        "tools": "All tools",
        "followup": [
            "Confirm each person's final presentation slot for Month 12.",
            "Send Month 12 briefing: presentation format, judging criteria.",
            "Coordinate leadership attendance for Grand Finale.",
        ],
    },
    {
        "num": 12, "title": "EMBRACE AI Showcase: Grand Finale", "phase": "DELIVER", "date": "March 10, 2026",
        "duration": "150 min", "type": "Presentations + Celebration (Extended session)",
        "seminar": "Our AI Journey: 12 Months of Growth — Retrospective (Vigneshvar SA)",
        "description": (
            "The culmination of the entire initiative. Each person presents their completed AI "
            "use case to the full team and leadership. Judging panel scores. Annual leaderboard "
            "finale with Star Points and celebrations."
        ),
        "prep": [
            "Collect final submissions: prototype, slides, impact metrics.",
            "Set up judging panel: Vigneshvar + Manager + 1–2 senior members or US stakeholder.",
            "Prepare judging scorecards (Business Relevance 30%, Technical Quality 25%, Creativity 20%, Presentation 15%, Reusability 10%).",
            "Prepare annual leaderboard with final totals.",
            "Arrange prizes: Star Points for Top 3.",
            "Book larger room + snacks/celebration items.",
        ],
        "execution": [
            "0:00–0:15 — Vigneshvar's retrospective: 12 Months of Growth.",
            "0:15–1:55 — Use Case Showcase: 5 min presentation + 2 min Q&A each (16 presentations).",
            "1:55–2:10 — Break / Judges deliberate.",
            "2:10–2:25 — Use Case Awards: Best Overall (15 pts), Runner-Up (10 pts), Second Runner-Up (7 pts), plus special categories.",
            "2:25–2:40 — Annual Leaderboard Finale. Star Points to Top 3. Certificates for all.",
            "2:40–2:50 — Manager remarks and next steps.",
        ],
        "awards": ["Best Overall AI Use Case (15 pts)", "Runner-Up (10 pts)", "Second Runner-Up (7 pts)", "Most Innovative Solution", "Most Business Impact", "Best Presentation", "People's Choice"],
        "skills": "End-to-end AI value demonstration, presentation skills",
        "tools": "All tools",
        "followup": [
            "Compile EMBRACE AI Portfolio: 16 use cases with executive summaries.",
            "Share portfolio with higher management in India and US profit centers.",
            "Process Star Points for Top 3.",
            "Archive all materials to shared repository.",
            "Propose operationalizing top 3–5 use cases. Outline Year 2 plan.",
        ],
    },
]

for md in months_data:
    doc.add_heading(f"9.{md['num']}  Month {md['num']} — {md['title']}", level=2)
    
    add_table(
        ["Attribute", "Details"],
        [
            ["Phase", md['phase']],
            ["Date", md['date']],
            ["Duration", md['duration']],
            ["Activity Type", md['type']],
            ["Seminar Topic", md['seminar']],
        ],
    )

    doc.add_heading("Description", level=3)
    doc.add_paragraph(md['description'])

    doc.add_heading("Preparation Steps", level=3)
    for step in md['prep']:
        p = doc.add_paragraph(step, style='List Bullet')

    doc.add_heading("Session Execution (Run of Show)", level=3)
    for step in md['execution']:
        p = doc.add_paragraph(step, style='List Bullet')

    doc.add_heading("Award Categories", level=3)
    for a in md['awards']:
        doc.add_paragraph(f"⭐ {a}", style='List Bullet')

    doc.add_heading("Skills Developed", level=3)
    doc.add_paragraph(md['skills'])

    doc.add_heading("Tools", level=3)
    doc.add_paragraph(md['tools'])

    doc.add_heading("Follow-Up Actions", level=3)
    for f in md['followup']:
        doc.add_paragraph(f, style='List Bullet')

    if md['num'] < 12:
        doc.add_page_break()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 10. EMBRACE AI DASHBOARD
# ═══════════════════════════════════════════════════════════
doc.add_heading("10. EMBRACE AI Dashboard (Web Application)", level=1)

doc.add_paragraph(
    "A custom-built web application powers the initiative, providing real-time tracking, "
    "gamification, and collaboration features. The dashboard is accessible to all team members "
    "via a single URL without requiring login."
)

doc.add_heading("10.1 Architecture", level=2)
doc.add_paragraph(
    "Backend: Python (Flask + Flask-SocketIO) — handles API requests and real-time WebSocket communication.\n"
    "Frontend: Single-page application (HTML + CSS + JavaScript) with Siemens-themed design.\n"
    "Database: JSON file (db.json) — lightweight, no external database required.\n"
    "Real-time: Socket.IO for live quiz, live survey, and instant leaderboard updates."
)

doc.add_heading("10.2 Features — Public View (No Login)", level=2)
features_public = [
    ["Dashboard", "Live countdown to next session, current activity card, seminar spotlight with presenter info, top 3 leaderboard preview, phase progress tracker."],
    ["Leaderboard", "Full table with avatars, member names, domain tags, monthly point breakdown (M1–M12), and cumulative totals. Sorted by rank."],
    ["Calendar", "12-month event timeline with phase indicators, session details, seminar topics, and status badges (Upcoming / Completed)."],
    ["Team", "Card grid of all 16 members with profile photos, roles, domain badges, and domain filter (All / PLM / Manufacturing / OT & Automation / Cross-Functional)."],
    ["Quiz (Join)", "Participants enter a room code to join a live quiz session hosted by the admin."],
    ["Survey (Vote)", "When a survey is live, an anonymous 1–5 star rating form appears on everyone's dashboard."],
]
add_table(["Feature", "Description"], features_public)

doc.add_heading("10.3 Features — Admin View (Login Required)", level=2)
doc.add_paragraph("Login: Click profile icon → enter admin/admin.")
features_admin = [
    ["Manage Members", "Add, edit, remove team members. Upload profile photos. Change domains and roles."],
    ["Award Points", "Select month, select member, choose category, set point value. Supports custom categories."],
    ["Edit Events", "Change date, time, presenter, topic, description, status, comments for any event."],
    ["Quiz Manager", "Create 5-question MCQ quizzes with correct answers. Host live quiz rooms with auto-generated room codes."],
    ["Survey Manager", "Launch anonymous seminar rating surveys. See real-time vote count and average. Close surveys."],
]
add_table(["Feature", "Description"], features_admin)

doc.add_heading("10.4 Live Quiz — Anti-Cheat Features", level=2)
quiz_features = [
    ["Random Question Order", "Questions appear in different order for each participant."],
    ["Tab-Switch Detection", "If a participant switches browser tabs, a -1 point penalty is applied and the admin is notified in real-time."],
    ["Copy/Paste Disabled", "Ctrl+C, Ctrl+V, and right-click are blocked during quiz."],
    ["Live Admin Dashboard", "Admin sees each participant's score, tab switches, and completion status in real-time."],
]
add_table(["Feature", "Description"], quiz_features)

doc.add_heading("10.5 How to Run", level=2)
doc.add_paragraph("Prerequisites: Python 3.8+, pip.")
steps = [
    "Install dependencies: pip install -r requirements.txt",
    "Start the server: python app.py",
    "Open in browser: http://localhost:3000",
    "For network access: Share http://YOUR_IP:3000 with the team.",
    "For cloud hosting: Push to GitHub → deploy on Render/Railway (free).",
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ═══════════════════════════════════════════════════════════
# 11. EXPECTED OUTCOMES
# ═══════════════════════════════════════════════════════════
doc.add_heading("11. Expected Outcomes After 12 Months", level=1)

outcomes = [
    "All 16 members are comfortable using GenAI tools (SiemensGPT, M365 Copilot) in their daily work.",
    "16 documented AI use cases with working prototypes — ready to be pitched to US/Canada profit centers as value-adds.",
    "Measurable productivity gains — aggregated time savings across the team.",
    "A culture of AI-first thinking embedded into Engineering Systems.",
    "A reusable playbook that can be shared with other Siemens teams in India as a model for Embrace AI adoption.",
    "Direct alignment with the team's existing AI engagement under the OT & Automation pillar — turning a formal engagement into demonstrated capability.",
]
for o in outcomes:
    doc.add_paragraph(o, style='List Bullet')

# ═══════════════════════════════════════════════════════════
# 12. CUMULATIVE DELIVERABLES
# ═══════════════════════════════════════════════════════════
doc.add_heading("12. Cumulative Deliverables Produced Over 12 Months", level=1)

deliverables = [
    ["1", "AI Portraits Gallery", "Month 1"],
    ["2", "EngSys AI Comic Strips", "Month 2"],
    ["3", "EngSys Prompt Library", "Month 3"],
    ["4", "Copilot Cookbook for EngSys", "Month 4"],
    ["5", "Code with AI — Lessons Learned + Hallucination Patterns", "Month 5"],
    ["6", "Team Knowledge Assistant (prototype) + Doc Preparation Guide", "Month 6"],
    ["7", "Automation Catalog (16 automations)", "Month 7"],
    ["8", "Agent Design Patterns for EngSys", "Month 8"],
    ["9", "AI Data Analysis Findings", "Month 9"],
    ["10", "16 Approved Use Case Proposals", "Month 10"],
    ["11", "16 Working Prototypes (draft)", "Month 11"],
    ["12", "EMBRACE AI Portfolio — 16 Final Use Cases", "Month 12"],
]
add_table(["#", "Deliverable", "Created In"], deliverables)

# ═══════════════════════════════════════════════════════════
# 13. APPENDIX — EVENT CALENDAR
# ═══════════════════════════════════════════════════════════
doc.add_heading("13. Appendix — Event Calendar", level=1)

calendar = [
    ["1", "Apr 8, 2025", "Who's Who: AI Edition", "SPARK", "Pre-worked", "90 min"],
    ["2", "May 13, 2025", "AI Storyteller", "SPARK", "In-session", "90 min"],
    ["3", "Jun 10, 2025", "Prompt Battle Arena", "SPARK", "In-session", "100 min"],
    ["4", "Jul 8, 2025", "Copilot Power User", "BUILD", "Pre-worked + Demo", "100 min"],
    ["5", "Aug 12, 2025", "Code with AI", "BUILD", "Pre-worked", "100 min"],
    ["6", "Sep 9, 2025", "Smart Documentation / RAG", "BUILD", "Collaborative", "100 min"],
    ["7", "Oct 14, 2025", "Automate the Boring Stuff", "APPLY", "Pre-worked + Demo", "110 min"],
    ["8", "Nov 11, 2025", "AI Agents Hackathon", "APPLY", "In-session", "120 min"],
    ["9", "Dec 9, 2025", "Data Detective", "APPLY", "Pre-worked + Demo", "100 min"],
    ["10", "Jan 13, 2026", "Use Case: Problem Definition", "DELIVER", "Pre-worked + Workshop", "110 min"],
    ["11", "Feb 10, 2026", "Use Case: Build & Iterate", "DELIVER", "Peer Review", "120 min"],
    ["12", "Mar 10, 2026", "Grand Finale Showcase", "DELIVER", "Presentations", "150 min"],
]
add_table(["Month", "Date", "Title", "Phase", "Type", "Duration"], calendar)

# ─── Final note ───
doc.add_paragraph()
add_para("— End of Report —", italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para("vibe coded with AI ✦", italic=True, size=10, color=PETROL, align=WD_ALIGN_PARAGRAPH.CENTER)

# ═══ SAVE ═══
doc.save(OUT)
print(f"\n✅ Report saved to:\n   {OUT}\n")
